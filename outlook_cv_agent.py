import argparse
import re
import shutil
from datetime import datetime
from pathlib import Path

import pdfplumber
import win32com.client
from docx import Document
from openpyxl import Workbook, load_workbook


DEFAULT_FOLDER = "Inbox/CVs To Process"
DEFAULT_OUTPUT = "candidate_details.xlsx"
DEFAULT_ATTACHMENTS_DIR = "downloaded_cvs"
SHEET_NAME = "Candidates"
HEADERS = [
    "Candidate Name",
    "Email",
    "Contact Number",
    "Experience",
    "Current Company",
    "Skills",
    "CV File",
    "Email Subject",
    "Sender",
    "Received Time",
    "Processed At",
]

EMAIL_PATTERN = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_PATTERN = re.compile(r"(?:\+91[\s-]?)?(?:0[\s-]?)?[6-9]\d{2}[\s-]?\d{3}[\s-]?\d{4}\b")
EXPERIENCE_PATTERN = re.compile(
    r"(?:(?:total\s+)?(?:experience|exp)\s*[:\-]?\s*)?(\d+(?:\.\d+)?)\s*\+?\s*(?:years|year|yrs|yr)\b",
    re.IGNORECASE,
)

KNOWN_SKILLS = [
    "Oracle Fusion",
    "Oracle Fusion HCM",
    "Oracle Fusion Financials",
    "Oracle SCM",
    "Oracle EBS",
    "SQL",
    "PL/SQL",
    "OTBI",
    "BI Publisher",
    "OIC",
    "REST API",
    "SOAP",
    "Python",
    "Java",
    "Excel",
    "ERP",
    "HCM",
    "Finance",
    "Procurement",
    "SCM",
]


def ensure_workbook(path: Path):
    if path.exists():
        workbook = load_workbook(path)
        if SHEET_NAME not in workbook.sheetnames:
            sheet = workbook.create_sheet(SHEET_NAME)
            sheet.append(HEADERS)
            workbook.save(path)
        return workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = SHEET_NAME
    sheet.append(HEADERS)
    workbook.save(path)
    return workbook


def existing_cv_files(sheet) -> set[str]:
    headers = {str(cell.value): index for index, cell in enumerate(sheet[1], start=1) if cell.value}
    cv_column = headers.get("CV File")
    if not cv_column:
        return set()
    values = set()
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if len(row) >= cv_column and row[cv_column - 1]:
            values.add(str(row[cv_column - 1]).lower())
    return values


def get_outlook_folder(folder_path: str):
    outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
    parts = [part.strip() for part in folder_path.replace("\\", "/").split("/") if part.strip()]
    if not parts:
        raise ValueError("Folder path is empty.")

    folder = outlook.GetDefaultFolder(6) if parts[0].lower() == "inbox" else outlook.Folders.Item(parts[0])
    for part in parts[1:]:
        folder = folder.Folders.Item(part)
    return folder


def safe_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "_", name)
    return cleaned.strip() or "attachment"


def save_cv_attachments(folder, attachments_dir: Path, limit: int):
    attachments_dir.mkdir(parents=True, exist_ok=True)
    items = folder.Items
    items.Sort("[ReceivedTime]", True)
    saved = []
    processed_emails = 0

    for message in items:
        if processed_emails >= limit:
            break

        try:
            if message.Class != 43:
                continue
            processed_emails += 1

            for attachment in message.Attachments:
                filename = safe_filename(str(attachment.FileName))
                suffix = Path(filename).suffix.lower()
                if suffix not in {".pdf", ".docx"}:
                    continue

                timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
                output_file = attachments_dir / f"{timestamp}_{filename}"
                attachment.SaveAsFile(str(output_file))
                saved.append(
                    {
                        "path": output_file,
                        "subject": str(message.Subject or ""),
                        "sender": str(message.SenderName or ""),
                        "received_time": str(message.ReceivedTime),
                    }
                )
        except Exception as exc:
            print(f"Skipped one email because Outlook could not read it: {exc}")

    return saved


def extract_pdf_text(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    return ""


def first_non_empty_line(text: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if EMAIL_PATTERN.search(line) or PHONE_PATTERN.search(line):
            continue
        if len(line.split()) <= 6:
            return line
    return ""


def extract_experience(text: str) -> str:
    matches = EXPERIENCE_PATTERN.findall(text)
    if not matches:
        return ""
    numbers = [float(match) for match in matches]
    best = max(numbers)
    return f"{best:g} years"


def extract_current_company(text: str) -> str:
    patterns = [
        r"(?:current\s+company|current\s+employer|present\s+company|present\s+employer)\s*[:\-]\s*(.+)",
        r"(?:working\s+with|working\s+at)\s+(.+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            return value.splitlines()[0].strip(" .,-")
    return ""


def extract_skills(text: str) -> str:
    found = []
    lower_text = text.lower()
    for skill in KNOWN_SKILLS:
        if skill.lower() in lower_text:
            found.append(skill)
    return ", ".join(dict.fromkeys(found))


def parse_candidate_details(path: Path):
    text = extract_text(path)
    emails = sorted(set(EMAIL_PATTERN.findall(text)))
    phones = sorted(set(match.group(0).strip() for match in PHONE_PATTERN.finditer(text)))
    return {
        "name": first_non_empty_line(text),
        "email": ", ".join(emails),
        "phone": ", ".join(phones),
        "experience": extract_experience(text),
        "company": extract_current_company(text),
        "skills": extract_skills(text),
    }


def run_agent(args):
    output_path = Path(args.output).resolve()
    attachments_dir = Path(args.attachments_dir).resolve()
    folder = get_outlook_folder(args.folder)
    attachments = save_cv_attachments(folder, attachments_dir, args.email_limit)

    workbook = ensure_workbook(output_path)
    sheet = workbook[SHEET_NAME]
    seen_files = existing_cv_files(sheet)
    processed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    added = 0

    for attachment in attachments:
        cv_path = attachment["path"]
        if str(cv_path).lower() in seen_files:
            continue

        try:
            details = parse_candidate_details(cv_path)
            sheet.append(
                [
                    details["name"],
                    details["email"],
                    details["phone"],
                    details["experience"],
                    details["company"],
                    details["skills"],
                    str(cv_path),
                    attachment["subject"],
                    attachment["sender"],
                    attachment["received_time"],
                    processed_at,
                ]
            )
            added += 1
        except Exception as exc:
            print(f"Could not process {cv_path.name}: {exc}")

    workbook.save(output_path)

    if args.clean_attachments:
        shutil.rmtree(attachments_dir, ignore_errors=True)

    print("\nDone.")
    print(f"CV attachments found: {len(attachments)}")
    print(f"New candidates added: {added}")
    print(f"Excel file: {output_path}")


def build_parser():
    parser = argparse.ArgumentParser(description="Read CV attachments from Outlook and create candidate Excel data.")
    parser.add_argument("--folder", default=DEFAULT_FOLDER, help="Outlook folder path, for example Inbox/CVs To Process.")
    parser.add_argument("--output", default=DEFAULT_OUTPUT, help="Excel output file.")
    parser.add_argument("--attachments-dir", default=DEFAULT_ATTACHMENTS_DIR, help="Folder for saved CV attachments.")
    parser.add_argument("--email-limit", type=int, default=50, help="Maximum recent emails to scan from the Outlook folder.")
    parser.add_argument("--clean-attachments", action="store_true", help="Delete downloaded CV files after Excel is created.")
    return parser


if __name__ == "__main__":
    run_agent(build_parser().parse_args())
