import argparse
import email
import getpass
import hashlib
import imaplib
import json
import logging
import os
import re
import shutil
import tempfile
import warnings
from datetime import datetime
from email.header import decode_header
from email.utils import parseaddr
from email.utils import parsedate_to_datetime
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

warnings.filterwarnings("ignore")
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfminer.pdfpage").setLevel(logging.ERROR)
logging.getLogger("pdfminer.pdfinterp").setLevel(logging.ERROR)
logging.getLogger("pdfminer.converter").setLevel(logging.ERROR)

try:
    import fitz
    import pytesseract
    from PIL import Image
except ImportError:
    fitz = None
    pytesseract = None
    Image = None


DEFAULT_OUTPUT = "Hiring_Automation_Candidates.xlsx"
DEFAULT_ATTACHMENTS_DIR = "downloaded_cvs"
DEFAULT_STATE_FILE = "imap_cv_agent_state.json"
SHEET_NAME = "Candidates"
HEADERS = [
    "Candidate Name",
    "Email",
    "Contact Number",
    "Experience",
    "Current Company",
    "Skills",
    "CV File",
    "Email Subject",
    "Sender",
    "Received Time",
    "Processed At",
    "Resume Hash",
]

EMAIL_PATTERN = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_CANDIDATE_PATTERN = re.compile(r"(?:\+?\d{1,3}[\s().-]*)?(?:0[\s().-]*)?\d[\d\s().-]{8,}\d")
EXPERIENCE_PATTERN = re.compile(
    r"(?:(?:total\s+)?(?:experience|exp)\s*[:\-]?\s*)?(\d+(?:\.\d+)?)\s*\+?\s*(?:years|year|yrs|yr)\b",
    re.IGNORECASE,
)
NAME_LABEL_PATTERN = re.compile(r"^(?:name|candidate\s+name)\s*[:\-]\s*(.+)$", re.IGNORECASE)
BAD_NAME_WORDS = {
    "resume",
    "resumé",
    "cv",
    "curriculum vitae",
    "profile",
    "biodata",
    "bio data",
    "personal details",
    "contact",
    "contact details",
}
BAD_NAME_PARTS = {
    "resume",
    "curriculum",
    "vitae",
    "profile",
    "email",
    "mail",
    "mobile",
    "phone",
    "contact",
    "address",
    "linkedin",
    "github",
    "objective",
    "summary",
    "experience",
    "education",
    "skills",
    "career",
    "overview",
    "acting",
    "lead",
    "designation",
    "qa",
    "sdet",
    "api",
    "automation",
    "testing",
    "linkedin",
    "www",
    "perspective",
    "mind",
    "clients",
    "services",
    "training",
    "gmail",
    "project",
    "accounting",
    "module",
    "bangalore",
    "india",
    "previous",
    "roles",
    "reliability",
    "streamline",
    "streamlinetestingprocesses",
    "andmentorjuniortesters",
    "whilecontinuouslyevolvingmytechnical",
    "python",
    "allinfo",
    "march",
    "april",
    "may",
    "june",
    "july",
    "august",
    "september",
    "october",
    "november",
    "december",
    "jan",
    "feb",
    "mar",
    "apr",
    "jun",
    "jul",
    "aug",
    "sep",
    "sept",
    "oct",
    "nov",
    "dec",
    "page",
    "only",
    "backend",
    "frontend",
    "manual",
    "mobile",
    "auto",
    "db",
    "ppm",
    "hcm",
    "financials",
    "finance",
    "compressed",
    "freelance",
    "functional",
    "associate",
    "in",
    "of",
    "sap",
    "abap",
    "fiori",
    "erp",
    "middleware",
    "osb",
    "btp",
    "cpi",
    "pipo",
    "date",
    "odata",
    "sftp",
    "https",
    "idoc",
    "jms",
    "integration",
    "adapters",
}
INTERNAL_EMAIL_DOMAINS = {"emotifzone.com"}
NON_CV_FILENAME_WORDS = {
    "payslip",
    "pay slip",
    "salary",
    "bill",
    "address proof",
    "residency",
    "visa",
    "passport",
    "aadhar",
    "aadhaar",
    "pan",
    "certificate",
    "certification",
    "certified",
    "implementation professional",
    "invoice",
    "photo",
    "image",
    "training",
    "invitation",
    "brochure",
    "company profile",
    "jd",
    "job description",
    "requirement",
}
CV_FILENAME_WORDS = {
    "resume",
    "cv",
    "curriculum",
    "profile",
}
CV_TEXT_WORDS = {
    "professional summary",
    "career objective",
    "work experience",
    "professional experience",
    "employment history",
    "technical skills",
    "core skills",
    "education",
}

KNOWN_SKILLS = [
    "Oracle Fusion",
    "Oracle Fusion HCM",
    "Oracle Fusion Financials",
    "Oracle SCM",
    "Oracle EBS",
    "SQL",
    "PL/SQL",
    "OTBI",
    "BI Publisher",
    "OIC",
    "REST API",
    "SOAP",
    "Python",
    "Java",
    "Excel",
    "ERP",
    "HCM",
    "Finance",
    "Procurement",
    "SCM",
]

TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
]

for tesseract_path in TESSERACT_PATHS:
    if pytesseract and Path(tesseract_path).exists():
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        break


def decode_text(value: str | None) -> str:
    if not value:
        return ""
    parts = []
    for content, charset in decode_header(value):
        if isinstance(content, bytes):
            for encoding in [charset, "utf-8", "latin-1"]:
                if not encoding:
                    continue
                try:
                    parts.append(content.decode(encoding, errors="ignore"))
                    break
                except LookupError:
                    continue
        else:
            parts.append(content)
    return "".join(parts)


def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[\x00-\x1f\x7f]", " ", decode_text(name))
    cleaned = re.sub(r'[<>:"/\\|?*]', "_", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip() or "attachment"


def ensure_workbook(path: Path):
    if path.exists():
        workbook = load_workbook(path)
        if SHEET_NAME not in workbook.sheetnames:
            sheet = workbook.create_sheet(SHEET_NAME)
            sheet.append(HEADERS)
            workbook.save(path)
        else:
            ensure_headers(workbook[SHEET_NAME])
            workbook.save(path)
        return workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = SHEET_NAME
    sheet.append(HEADERS)
    workbook.save(path)
    return workbook


def ensure_headers(sheet):
    existing = [cell.value for cell in sheet[1]]
    if existing == HEADERS:
        return

    rows = []
    for values in sheet.iter_rows(min_row=2, values_only=True):
        row_data = {}
        for index, header in enumerate(existing):
            if header:
                row_data[str(header)] = values[index] if index < len(values) else ""
        rows.append(row_data)

    if sheet.max_row:
        sheet.delete_rows(1, sheet.max_row)

    sheet.append(HEADERS)
    for row_data in rows:
        sheet.append([row_data.get(header, "") for header in HEADERS])


def header_map(sheet) -> dict[str, int]:
    return {str(cell.value): index for index, cell in enumerate(sheet[1], start=1) if cell.value}


def normalize_email(value: str) -> str:
    return value.strip().lower()


def is_internal_email(value: str) -> bool:
    email_value = normalize_email(value)
    if "@" not in email_value:
        return False
    return email_value.rsplit("@", 1)[1] in INTERNAL_EMAIL_DOMAINS


def normalize_name(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip().lower()


def normalize_phone(value: str) -> str:
    digits = re.sub(r"\D", "", value)
    if len(digits) > 10 and digits[-10] in "6789":
        return digits[-10:]
    if 10 <= len(digits) <= 15 and not digits.startswith("91"):
        return digits
    if len(digits) > 10 and digits.startswith("91"):
        digits = digits[-10:]
    if len(digits) == 11 and digits.startswith("0"):
        digits = digits[1:]
    return digits if len(digits) == 10 and digits[0] in "6789" else ""


def candidate_keys(details: dict[str, str]) -> set[str]:
    keys = set()
    name_value = normalize_name(details.get("name", ""))
    if name_value:
        keys.add(f"name:{name_value}")

    for email_value in details.get("email", "").split(","):
        email_value = normalize_email(email_value)
        if email_value and not is_internal_email(email_value):
            keys.add(f"email:{email_value}")

    for phone_value in details.get("phone", "").split(","):
        phone_value = normalize_phone(phone_value)
        if phone_value:
            keys.add(f"phone:{phone_value}")
    return keys


def is_usable_candidate(details: dict[str, str]) -> bool:
    has_name = bool(normalize_name(details.get("name", "")))
    has_email = any(normalize_email(value) and not is_internal_email(value) for value in details.get("email", "").split(","))
    has_phone = any(normalize_phone(value) for value in details.get("phone", "").split(","))
    return has_name or has_email or has_phone


def existing_candidate_rows(sheet) -> dict[str, int]:
    headers = header_map(sheet)
    keys_to_rows = {}

    for row_number in range(2, sheet.max_row + 1):
        details = {
            "name": str(sheet.cell(row=row_number, column=headers.get("Candidate Name", 0)).value or ""),
            "email": str(sheet.cell(row=row_number, column=headers.get("Email", 0)).value or ""),
            "phone": str(sheet.cell(row=row_number, column=headers.get("Contact Number", 0)).value or ""),
        }

        for key in candidate_keys(details):
            keys_to_rows[key] = row_number

    return keys_to_rows


def connect_imap(server: str, port: int, username: str, password: str):
    mailbox = imaplib.IMAP4_SSL(server, port)
    mailbox.login(username, password)
    return mailbox


def load_state(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def save_state(path: Path, state: dict):
    path.write_text(json.dumps(state, indent=2), encoding="utf-8")


def state_key(server: str, username: str, folder: str) -> str:
    return f"{server.lower()}|{username.lower()}|{folder}"


def find_legacy_checkpoint(state: dict, server: str, username: str, folder: str) -> int | None:
    matches = []
    for value in state.values():
        if not isinstance(value, dict):
            continue
        if (
            str(value.get("server", "")).lower() == server.lower()
            and str(value.get("username", "")).lower() == username.lower()
            and str(value.get("folder", "")) == folder
            and value.get("last_processed_uid") is not None
        ):
            matches.append(int(value["last_processed_uid"]))
    return max(matches) if matches else None


def get_message_uids(mailbox, email_limit: int, last_processed_uid: int | None):
    status, data = mailbox.uid("search", None, "ALL")
    if status != "OK":
        raise RuntimeError("Could not search mailbox.")

    all_uids = [int(uid) for uid in data[0].split()]
    if last_processed_uid is not None:
        return [uid for uid in all_uids if uid > last_processed_uid]

    return all_uids[-email_limit:]


def get_highest_uid(mailbox, folder: str) -> int:
    status, _ = mailbox.select(f'"{folder}"')
    if status != "OK":
        raise RuntimeError(f"Could not open mail folder: {folder}")
    status, data = mailbox.uid("search", None, "ALL")
    if status != "OK":
        raise RuntimeError("Could not search mailbox.")
    all_uids = [int(uid) for uid in data[0].split()]
    return max(all_uids) if all_uids else 0


def save_cv_attachments(mailbox, folder: str, attachments_dir: Path, email_limit: int, last_processed_uid: int | None):
    attachments_dir.mkdir(parents=True, exist_ok=True)
    status, _ = mailbox.select(f'"{folder}"')
    if status != "OK":
        raise RuntimeError(f"Could not open mail folder: {folder}")

    message_uids = get_message_uids(mailbox, email_limit, last_processed_uid)
    saved = []
    highest_uid = last_processed_uid or 0

    for message_uid in message_uids:
        highest_uid = max(highest_uid, message_uid)
        status, data = mailbox.uid("fetch", str(message_uid), "(RFC822)")
        if status != "OK" or not data or not data[0]:
            continue

        message = email.message_from_bytes(data[0][1])
        subject = decode_text(message.get("Subject"))
        sender = decode_text(message.get("From"))
        sender_name, sender_email = parse_sender(sender)
        received_time = decode_text(message.get("Date"))

        for part in message.walk():
            if part.get_content_disposition() != "attachment":
                continue

            filename = safe_filename(part.get_filename())
            suffix = Path(filename).suffix.lower()
            if suffix not in {".pdf", ".docx"}:
                continue

            payload = part.get_payload(decode=True)
            if not payload:
                continue

            temp_file = save_payload_to_temp_file(payload, filename)
            try:
                details = parse_candidate_details(temp_file, sender_name=sender_name, sender_email=sender_email)
            except Exception as exc:
                print(f"Skipped unreadable attachment: {temp_file.name} ({exc})")
                temp_file.unlink(missing_ok=True)
                continue
            if not details.get("is_cv"):
                print(f"Skipped non-CV attachment: {temp_file.name}")
                temp_file.unlink(missing_ok=True)
                continue

            output_file = move_unique_file(temp_file, attachments_dir)
            saved.append(
                {
                    "path": output_file,
                    "details": details,
                    "subject": subject,
                    "sender": sender,
                    "sender_name": sender_name,
                    "sender_email": sender_email,
                    "received_time": received_time,
                    "message_uid": message_uid,
                }
            )

    return saved, len(message_uids), highest_uid


def save_payload_to_temp_file(payload: bytes, filename: str) -> Path:
    temp_dir = Path(tempfile.gettempdir()) / "HiringAutomationTool"
    temp_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    temp_file = temp_dir / f"{timestamp}_{filename}"
    temp_file.write_bytes(payload)
    return temp_file


def move_unique_file(source: Path, destination_dir: Path) -> Path:
    destination_dir.mkdir(parents=True, exist_ok=True)
    destination = destination_dir / source.name
    if not destination.exists():
        shutil.move(str(source), str(destination))
        return destination

    stem = destination.stem
    suffix = destination.suffix
    counter = 2
    while True:
        candidate = destination_dir / f"{stem}_{counter}{suffix}"
        if not candidate.exists():
            shutil.move(str(source), str(candidate))
            return candidate
        counter += 1


def extract_pdf_text(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    text = "\n".join(text_parts).strip()
    if len(text) >= 40:
        return text
    return extract_pdf_text_with_ocr(path)


def extract_pdf_text_with_ocr(path: Path) -> str:
    if not fitz or not pytesseract or not Image:
        print(f"OCR not available. Could not read scanned PDF: {path.name}")
        return ""

    text_parts = []
    try:
        document = fitz.open(path)
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            text_parts.append(pytesseract.image_to_string(image))
    except Exception as exc:
        print(f"OCR failed for {path.name}: {exc}")
        return ""
    return "\n".join(text_parts).strip()


def extract_docx_text(path: Path) -> str:
    try:
        document = Document(path)
    except Exception as exc:
        print(f"Could not read DOCX attachment {path.name}: {exc}")
        return ""
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def extract_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        return extract_pdf_text(path)
    if path.suffix.lower() == ".docx":
        return extract_docx_text(path)
    return ""


def first_non_empty_line(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    for line in lines[:25]:
        match = NAME_LABEL_PATTERN.match(line)
        if match:
            candidate = clean_name(match.group(1))
            if is_likely_name(candidate):
                return candidate

    contact_line_index = next(
        (index for index, line in enumerate(lines[:30]) if EMAIL_PATTERN.search(line) or extract_phone_numbers(line)),
        None,
    )
    if contact_line_index is not None:
        for index in range(max(0, contact_line_index - 4), contact_line_index):
            candidate = clean_name(lines[index])
            if is_likely_name(candidate):
                return candidate

    for line in lines[:40]:
        candidate = clean_name(line)
        if is_likely_name(candidate):
            return candidate
    return ""


def clean_name(value: str) -> str:
    value = re.sub(r"\b(resume|cv)\b", "", value, flags=re.IGNORECASE)
    value = re.sub(r"[^A-Za-z .'-]", " ", value)
    value = re.sub(r"\s+", " ", value).strip(" .-")
    value = re.sub(r"\s+[A-Za-z]$", "", value).strip(" .-")
    return value


def is_likely_name(value: str, allow_job_words: bool = False, allow_single_word: bool = False) -> bool:
    if not value:
        return False
    normalized = normalize_name(value)
    if normalized in BAD_NAME_WORDS:
        return False
    parts = value.split()
    min_words = 1 if allow_single_word else 2
    if not min_words <= len(parts) <= 5:
        return False
    if not allow_job_words and any(part.lower().strip(".:-") in BAD_NAME_PARTS for part in parts):
        return False
    if EMAIL_PATTERN.search(value) or extract_phone_numbers(value):
        return False
    letters = re.sub(r"[^A-Za-z]", "", value)
    if len(letters) < 4:
        return False
    return True


def extract_experience(text: str) -> str:
    matches = EXPERIENCE_PATTERN.findall(text)
    if not matches:
        return ""
    numbers = [float(match) for match in matches if 0 <= float(match) <= 40]
    if not numbers:
        return ""
    return f"{max(numbers):g} years"


def is_unrealistic_experience(value: object) -> bool:
    if not value:
        return False
    match = re.search(r"(\d+(?:\.\d+)?)", str(value))
    return bool(match and float(match.group(1)) > 40)


def extract_current_company(text: str) -> str:
    patterns = [
        r"(?:current\s+company|current\s+employer|present\s+company|present\s+employer)\s*[:\-]\s*(.+)",
        r"(?:working\s+with|working\s+at)\s+(.+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).splitlines()[0].strip(" .,-")
    return ""


def extract_skills(text: str) -> str:
    lower_text = text.lower()
    found = [skill for skill in KNOWN_SKILLS if skill.lower() in lower_text]
    return ", ".join(dict.fromkeys(found))


def extract_phone_numbers(text: str) -> list[str]:
    phones = []
    for match in PHONE_CANDIDATE_PATTERN.finditer(text):
        phone = normalize_phone(match.group(0))
        if phone:
            phones.append(phone)
    return sorted(set(phones))


def extract_emails(text: str) -> list[str]:
    return sorted({email_value for email_value in EMAIL_PATTERN.findall(text) if not is_internal_email(email_value)})


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def parse_sender(value: str) -> tuple[str, str]:
    display_name, email_address = parseaddr(value or "")
    return decode_text(display_name).strip().strip('"'), email_address


def parse_candidate_details(path: Path, sender_name: str = "", sender_email: str = ""):
    text = extract_text(path)
    if not is_likely_cv_document(path, text):
        return {
            "name": "",
            "email": "",
            "phone": "",
            "experience": "",
            "company": "",
            "skills": "",
            "is_cv": False,
        }

    emails = extract_emails(text)
    phones = extract_phone_numbers(text)
    name = choose_candidate_name(text, path, sender_name, sender_email)
    return {
        "name": name,
        "email": ", ".join(emails),
        "phone": ", ".join(phones),
        "experience": extract_experience(text),
        "company": extract_current_company(text),
        "skills": extract_skills(text),
        "is_cv": True,
    }


def is_likely_cv_document(path: Path, text: str) -> bool:
    filename = normalize_name(path.name)
    if is_non_cv_filename(path):
        return False
    if any(word in filename for word in CV_FILENAME_WORDS):
        return True

    lower_text = text.lower()
    if any(word in lower_text for word in CV_TEXT_WORDS):
        return True

    has_contact = bool(extract_emails(text) or extract_phone_numbers(text))
    has_experience = bool(EXPERIENCE_PATTERN.search(text))
    has_skills = bool(extract_skills(text))
    return has_contact and (has_experience or has_skills)


def is_non_cv_filename(path: Path) -> bool:
    filename = normalize_name(path.name)
    return any(word in filename for word in NON_CV_FILENAME_WORDS)


def needs_row_refresh(row_data: dict[str, object]) -> bool:
    current_name = str(row_data.get("Candidate Name") or "")
    if not current_name or is_weak_name(current_name):
        return True
    if len(clean_name(current_name).split()) < 2:
        return True
    if not row_data.get("Email") and not row_data.get("Contact Number"):
        return True
    if is_unrealistic_experience(row_data.get("Experience")):
        return True
    return False


def choose_candidate_name(text: str, path: Path, sender_name: str = "", sender_email: str = "") -> str:
    file_name = name_from_filename(path)
    sender_candidate = name_from_sender(sender_name, sender_email)
    text_name = first_non_empty_line(text)
    if sender_candidate and (not file_name or is_weak_name(file_name) or has_name_noise(file_name)):
        return sender_candidate
    strong_candidates = [
        candidate for candidate in [file_name, sender_candidate] if len(candidate.split()) >= 2
    ]
    candidates = strong_candidates + [file_name, sender_candidate, text_name]
    for candidate in candidates:
        if candidate and not is_weak_name(candidate):
            return candidate
    for candidate in candidates:
        if candidate:
            return candidate
    return ""


def is_weak_name(value: str) -> bool:
    if not value:
        return True
    cleaned = clean_name(value)
    if cleaned and cleaned != value:
        return is_weak_name(cleaned)
    parts = {re.sub(r"[^a-z]", "", part.lower()) for part in value.split()}
    compact = re.sub(r"[^a-z]", "", value.lower())
    if compact in BAD_NAME_WORDS or compact in BAD_NAME_PARTS:
        return True
    weak_parts = BAD_NAME_PARTS | {
        "oracle",
        "fusion",
        "clients",
        "client",
        "consultant",
        "functional",
        "technical",
        "senior",
        "engineer",
        "developer",
        "qa",
        "sdet",
        "testing",
        "automation",
        "gmail",
        "linkedin",
        "career",
        "overview",
        "designation",
        "acting",
        "lead",
        "services",
        "training",
        "program",
        "of",
        "sap",
        "abap",
        "fiori",
        "erp",
        "middleware",
        "osb",
        "btp",
    }
    if any(marker in compact for marker in {"gmail", "linkedin", "email", "www"}):
        return True
    return bool(parts & weak_parts)


def has_name_noise(value: str) -> bool:
    parts = {re.sub(r"[^a-z]", "", part.lower()) for part in value.split()}
    noise = {
        "python",
        "allinfo",
        "march",
        "december",
        "page",
        "only",
        "backend",
        "engineer",
        "finance",
        "financials",
        "ppm",
        "hcm",
        "compressed",
        "freelance",
        "functional",
        "manual",
        "mobile",
        "auto",
        "db",
        "associate",
        "in",
        "of",
        "sap",
        "abap",
        "fiori",
        "erp",
        "middleware",
        "osb",
        "btp",
        "cpi",
        "pipo",
        "date",
        "odata",
        "sftp",
        "https",
        "idoc",
        "jms",
        "integration",
        "adapters",
    }
    return bool(parts & noise)


def should_replace_name(current_name: str, new_name: str) -> bool:
    if not new_name:
        return False
    if not current_name:
        return True
    if (is_weak_name(current_name) or has_name_noise(current_name)) and not is_weak_name(new_name):
        return True
    return len(new_name.split()) >= 2 and len(clean_name(current_name).split()) < 2


def split_compact_name(value: str) -> str:
    value = re.sub(r"([a-z])([A-Z])", r"\1 \2", value or "")
    value = re.sub(r"\b([A-Za-z]{3,})x([A-Za-z]{2,})\b", r"\1 \2", value)
    value = re.sub(r"([A-Za-z])(\d)", r"\1 \2", value)
    value = re.sub(r"(\d)([A-Za-z])", r"\1 \2", value)
    return value


def name_from_sender(sender_name: str, sender_email: str = "") -> str:
    source = sender_name
    if not source and sender_email and not is_internal_email(sender_email):
        source = sender_email.split("@", 1)[0]
    source = split_compact_name(source)
    source = re.sub(r"[_+.-]+", " ", source)
    source = clean_name(source)
    source = re.sub(r"\b(gmail|yahoo|outlook|hotmail|email|mail)\b", " ", source, flags=re.IGNORECASE)
    words = [word for word in source.split() if not re.search(r"\d", word) and len(word) > 1]
    candidate = clean_name(" ".join(word.capitalize() for word in words[:5]))
    return candidate if is_likely_name(candidate, allow_job_words=True, allow_single_word=True) else ""


def name_from_filename(path: Path) -> str:
    name = re.sub(r"^\d{14,20}_", "", path.stem)
    name = re.sub(r"\.(pdf|docx?)$", "", name, flags=re.IGNORECASE)
    name = split_compact_name(name)
    name = re.sub(r"[_+.-]+", " ", name)
    name = re.sub(
        r"\b(resume|resumee|cv|curriculum|vitae|updated|latest|final|copy|new|profile|ats|full|pdf|docx?|oracle|fusion|consultant|functional|technical|python|allinfo|march|april|may|june|july|august|september|october|november|december|page|only|backend|engineer|finance|financials|ppm|hcm|compressed|freelance|manual|mobile|auto|db|associate|in|of|sap|abap|fiori|erp|middleware|osb|btp|cpi|pipo|date|odata|sftp|https|idoc|jms|integration|adapters)\b",
        " ",
        name,
        flags=re.IGNORECASE,
    )
    name = re.sub(r"\b\d+(?:\.\d+)?\s*(?:yrs?|years?)?\b", " ", name, flags=re.IGNORECASE)
    name = re.sub(r"\s+", " ", name).strip(" .-_")
    words = [word for word in name.split() if len(word) > 1]
    role_words = {
        "senior",
        "oracle",
        "fusion",
        "ebs",
        "finance",
        "financial",
        "scm",
        "hcm",
        "functional",
        "technical",
        "consultant",
        "developer",
        "engineer",
        "python",
        "allinfo",
        "march",
        "december",
        "page",
        "only",
        "backend",
        "ppm",
        "freelance",
        "compressed",
        "in",
        "of",
        "sap",
        "abap",
        "fiori",
        "erp",
        "middleware",
        "osb",
        "btp",
        "cpi",
        "pipo",
        "date",
    }
    trimmed_words = []
    for word in words:
        if word.lower() in role_words and len(trimmed_words) >= 2:
            break
        trimmed_words.append(word)
    words = trimmed_words[:5]
    candidate = clean_name(" ".join(word.capitalize() for word in words))
    return candidate if is_likely_name(candidate, allow_job_words=True, allow_single_word=True) else ""


def parse_excel_date(value):
    if isinstance(value, datetime):
        return value
    if not value:
        return datetime.min
    text = str(value).strip()
    for parser in (
        lambda item: parsedate_to_datetime(item).replace(tzinfo=None),
        lambda item: datetime.fromisoformat(item.replace("Z", "+00:00")).replace(tzinfo=None),
    ):
        try:
            return parser(text)
        except Exception:
            continue
    for date_format in ("%Y-%m-%d %H:%M:%S", "%d-%m-%Y %H:%M:%S", "%d/%m/%Y %H:%M:%S"):
        try:
            return datetime.strptime(text, date_format)
        except ValueError:
            continue
    return datetime.min


def row_details(row_data: dict[str, object]) -> dict[str, str]:
    return {
        "name": str(row_data.get("Candidate Name") or ""),
        "email": str(row_data.get("Email") or ""),
        "phone": str(row_data.get("Contact Number") or ""),
    }


def clean_sort_and_format_sheet(sheet):
    headers = header_map(sheet)
    rows = []
    for values in sheet.iter_rows(min_row=2, values_only=True):
        row_data = {}
        for header in HEADERS:
            column = headers.get(header)
            row_data[header] = values[column - 1] if column and column - 1 < len(values) else ""
        cv_file = Path(str(row_data.get("CV File") or ""))
        if cv_file.exists():
            try:
                if is_non_cv_filename(cv_file):
                    continue
                if needs_row_refresh(row_data):
                    sender_name, sender_email = parse_sender(str(row_data.get("Sender") or ""))
                    details = parse_candidate_details(cv_file, sender_name=sender_name, sender_email=sender_email)
                    if not details.get("is_cv"):
                        continue
                    if should_replace_name(str(row_data.get("Candidate Name") or ""), details.get("name", "")):
                        row_data["Candidate Name"] = details["name"]
                    for detail_key, header in {
                        "email": "Email",
                        "phone": "Contact Number",
                        "experience": "Experience",
                        "company": "Current Company",
                        "skills": "Skills",
                    }.items():
                        should_fill = not row_data.get(header)
                        if header == "Experience" and is_unrealistic_experience(row_data.get(header)):
                            should_fill = True
                        if details.get(detail_key) and should_fill:
                            row_data[header] = details[detail_key]
            except Exception:
                pass
        if is_usable_candidate(row_details(row_data)):
            rows.append(row_data)

    rows.sort(
        key=lambda row: (
            parse_excel_date(row.get("Received Time")),
            parse_excel_date(row.get("Processed At")),
        ),
        reverse=True,
    )

    unique_rows = []
    seen = {}
    for row_data in rows:
        keys = candidate_keys(row_details(row_data))
        matching_index = next((seen[key] for key in keys if key in seen), None)
        if matching_index is not None:
            kept = unique_rows[matching_index]
            for header in HEADERS:
                if not kept.get(header) and row_data.get(header):
                    kept[header] = row_data[header]
                elif header == "Candidate Name" and should_replace_name(str(kept.get(header) or ""), str(row_data.get(header) or "")):
                    kept[header] = row_data[header]
                elif header == "Experience" and is_unrealistic_experience(kept.get(header)) and row_data.get(header):
                    kept[header] = row_data[header]
            continue

        seen_index = len(unique_rows)
        unique_rows.append(row_data)
        for key in keys:
            seen[key] = seen_index

    if sheet.max_row:
        sheet.delete_rows(1, sheet.max_row)

    sheet.append(HEADERS)
    for row_data in unique_rows:
        sheet.append([row_data.get(header, "") for header in HEADERS])

    header_fill = PatternFill(fill_type="solid", fgColor="D9EAF7")
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions

    for column_number, header in enumerate(HEADERS, start=1):
        letter = get_column_letter(column_number)
        if header in {"Email Subject", "Skills", "CV File"}:
            sheet.column_dimensions[letter].width = 35
        elif header in {"Received Time", "Processed At"}:
            sheet.column_dimensions[letter].width = 24
        else:
            sheet.column_dimensions[letter].width = max(14, min(28, len(header) + 4))

    if "Resume Hash" in header_map(sheet):
        sheet.column_dimensions[get_column_letter(header_map(sheet)["Resume Hash"])].hidden = True


def run_agent(args):
    password = args.password or os.environ.get("CV_AGENT_PASSWORD") or getpass.getpass("Enter email password or app password: ")
    state_path = Path(args.state_file).resolve()
    state = load_state(state_path)
    mailbox_state_key = state_key(args.imap_server, args.username, args.folder)
    last_processed_uid = None if args.full_scan else state.get(mailbox_state_key, {}).get("last_processed_uid")
    if last_processed_uid is None and not args.full_scan:
        last_processed_uid = find_legacy_checkpoint(state, args.imap_server, args.username, args.folder)
        if last_processed_uid is not None:
            print(f"Recovered incremental checkpoint from old state: UID {last_processed_uid}.")

    if last_processed_uid:
        print(f"Incremental mode: checking emails newer than UID {last_processed_uid}.")
    else:
        print(f"Initial/full scan mode: scanning latest {args.email_limit} emails.")

    mailbox = connect_imap(args.imap_server, args.port, args.username, password)

    if args.mark_current:
        highest_uid = get_highest_uid(mailbox, args.folder)
        mailbox.logout()
        state[mailbox_state_key] = {
            "last_processed_uid": highest_uid,
            "server": args.imap_server,
            "username": args.username,
            "folder": args.folder,
            "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        save_state(state_path, state)
        print(f"Incremental checkpoint set to current latest UID {highest_uid}.")
        print("Next normal run will check only emails received after this point.")
        return

    attachments, emails_checked, highest_uid = save_cv_attachments(
        mailbox,
        args.folder,
        Path(args.attachments_dir).resolve(),
        args.email_limit,
        last_processed_uid,
    )
    mailbox.logout()

    workbook = ensure_workbook(Path(args.output).resolve())
    sheet = workbook[SHEET_NAME]
    headers = header_map(sheet)
    existing_rows = existing_candidate_rows(sheet)
    processed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    added = 0
    updated = 0

    for attachment in attachments:
        cv_path = attachment["path"]

        try:
            details = attachment.get("details") or parse_candidate_details(
                cv_path,
                sender_name=attachment.get("sender_name", ""),
                sender_email=attachment.get("sender_email", ""),
            )

            if not is_usable_candidate(details):
                print(f"Skipped incomplete CV: {cv_path.name}")
                continue

            resume_hash = file_hash(cv_path)
            matching_row = next((existing_rows[key] for key in candidate_keys(details) if key in existing_rows), None)

            if matching_row:
                for header, value in {
                    "Candidate Name": details["name"],
                    "Email": details["email"],
                    "Contact Number": details["phone"],
                    "Experience": details["experience"],
                    "Current Company": details["company"],
                    "Skills": details["skills"],
                    "Resume Hash": resume_hash,
                }.items():
                    column = headers.get(header)
                    existing_value = sheet.cell(row=matching_row, column=column).value if column else None
                    should_update = not existing_value
                    if header == "Candidate Name":
                        should_update = should_replace_name(str(existing_value or ""), value)
                    if column and value and should_update:
                        sheet.cell(row=matching_row, column=column, value=value)
                        updated += 1
                continue

            sheet.append([
                details["name"],
                details["email"],
                details["phone"],
                details["experience"],
                details["company"],
                details["skills"],
                str(cv_path),
                attachment["subject"],
                attachment["sender"],
                attachment["received_time"],
                processed_at,
                resume_hash,
            ])
            new_row = sheet.max_row
            for candidate_key in candidate_keys(details):
                existing_rows[candidate_key] = new_row
            added += 1
        except Exception as exc:
            print(f"Could not process {cv_path.name}: {exc}")

    clean_sort_and_format_sheet(sheet)
    workbook.save(Path(args.output).resolve())

    state[mailbox_state_key] = {
        "last_processed_uid": highest_uid,
        "server": args.imap_server,
        "username": args.username,
        "folder": args.folder,
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    save_state(state_path, state)

    if args.clean_attachments:
        shutil.rmtree(Path(args.attachments_dir).resolve(), ignore_errors=True)

    print("\nDone.")
    print(f"Emails checked this run: {emails_checked}")
    print(f"CV attachments found: {len(attachments)}")
    print(f"New candidates added: {added}")
    print(f"Missing fields updated: {updated}")
    print(f"Excel file: {Path(args.output).resolve()}")
    print(f"Incremental state saved: {state_path}")


def build_parser():
    parser = argparse.ArgumentParser(description="Read CV attachments using IMAP and create candidate Excel data.")
    parser.add_argument("--imap-server", required=True, help="IMAP server, for example imap.gmail.com or mail.example.com.")
    parser.add_argument("--username", required=True, help="Email address, for example careers@emotifzone.com.")
    parser.add_argument("--password", help="Email password or app password. If omitted, you will be prompted.")
    parser.add_argument("--port", type=int, default=993, help="IMAP SSL port. Usually 993.")
    parser.add_argument("--folder", default="INBOX", help="Mail folder to scan. Start with INBOX.")
    parser.add_argument("--output", default=DEFAULT_OUTPUT, help="Excel output file.")
    parser.add_argument("--attachments-dir", default=DEFAULT_ATTACHMENTS_DIR, help="Folder for saved CV attachments.")
    parser.add_argument("--email-limit", type=int, default=500, help="Maximum recent emails to scan.")
    parser.add_argument("--state-file", default=DEFAULT_STATE_FILE, help="File that stores the last processed email UID.")
    parser.add_argument("--full-scan", action="store_true", help="Ignore incremental state and scan the latest emails again.")
    parser.add_argument("--mark-current", action="store_true", help="Set checkpoint to the current latest email without processing old emails.")
    parser.add_argument("--clean-attachments", action="store_true", help="Delete downloaded CV files after Excel is created.")
    return parser


if __name__ == "__main__":
    run_agent(build_parser().parse_args())
