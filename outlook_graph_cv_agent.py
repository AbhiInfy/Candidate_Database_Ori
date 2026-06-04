import argparse
import base64
import json
import re
import shutil
from datetime import datetime
from pathlib import Path

import msal
import pdfplumber
import requests
from docx import Document
from openpyxl import Workbook, load_workbook


GRAPH_BASE_URL = "https://graph.microsoft.com/v1.0"
DEFAULT_FOLDER = "CVs To Process"
DEFAULT_OUTPUT = "candidate_details.xlsx"
DEFAULT_ATTACHMENTS_DIR = "downloaded_cvs"
TOKEN_CACHE_FILE = "ms_graph_token_cache.json"
SHEET_NAME = "Candidates"
SCOPES = ["Mail.Read"]

HEADERS = [
    "Candidate Name",
    "Email",
    "Contact Number",
    "Experience",
    "Current Company",
    "Skills",
    "CV File",
    "Email Subject",
    "Sender",
    "Received Time",
    "Processed At",
]

EMAIL_PATTERN = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_PATTERN = re.compile(r"(?:\+91[\s-]?)?(?:0[\s-]?)?[6-9]\d{2}[\s-]?\d{3}[\s-]?\d{4}\b")
EXPERIENCE_PATTERN = re.compile(
    r"(?:(?:total\s+)?(?:experience|exp)\s*[:\-]?\s*)?(\d+(?:\.\d+)?)\s*\+?\s*(?:years|year|yrs|yr)\b",
    re.IGNORECASE,
)

KNOWN_SKILLS = [
    "Oracle Fusion",
    "Oracle Fusion HCM",
    "Oracle Fusion Financials",
    "Oracle SCM",
    "Oracle EBS",
    "SQL",
    "PL/SQL",
    "OTBI",
    "BI Publisher",
    "OIC",
    "REST API",
    "SOAP",
    "Python",
    "Java",
    "Excel",
    "ERP",
    "HCM",
    "Finance",
    "Procurement",
    "SCM",
]


def ensure_workbook(path: Path):
    if path.exists():
        workbook = load_workbook(path)
        if SHEET_NAME not in workbook.sheetnames:
            sheet = workbook.create_sheet(SHEET_NAME)
            sheet.append(HEADERS)
            workbook.save(path)
        return workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = SHEET_NAME
    sheet.append(HEADERS)
    workbook.save(path)
    return workbook


def existing_cv_files(sheet) -> set[str]:
    headers = {str(cell.value): index for index, cell in enumerate(sheet[1], start=1) if cell.value}
    cv_column = headers.get("CV File")
    if not cv_column:
        return set()
    values = set()
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if len(row) >= cv_column and row[cv_column - 1]:
            values.add(str(row[cv_column - 1]).lower())
    return values


def load_token_cache(path: Path):
    cache = msal.SerializableTokenCache()
    if path.exists():
        cache.deserialize(path.read_text(encoding="utf-8"))
    return cache


def save_token_cache(cache, path: Path):
    if cache.has_state_changed:
        path.write_text(cache.serialize(), encoding="utf-8")


def get_access_token(client_id: str, tenant: str):
    cache_path = Path(TOKEN_CACHE_FILE)
    cache = load_token_cache(cache_path)
    app = msal.PublicClientApplication(
        client_id=client_id,
        authority=f"https://login.microsoftonline.com/{tenant}",
        token_cache=cache,
    )

    accounts = app.get_accounts()
    result = app.acquire_token_silent(SCOPES, account=accounts[0]) if accounts else None

    if not result:
        flow = app.initiate_device_flow(scopes=SCOPES)
        if "user_code" not in flow:
            raise RuntimeError(f"Could not create device login flow: {json.dumps(flow, indent=2)}")

        print("\nMicrosoft sign-in required one time.")
        print(flow["message"])
        result = app.acquire_token_by_device_flow(flow)

    save_token_cache(cache, cache_path)

    if "access_token" not in result:
        raise RuntimeError(f"Could not get Microsoft Graph token: {json.dumps(result, indent=2)}")
    return result["access_token"]


def graph_get(token: str, url: str, params: dict | None = None):
    response = requests.get(url, headers={"Authorization": f"Bearer {token}"}, params=params, timeout=60)
    if response.status_code >= 400:
        raise RuntimeError(f"Graph request failed {response.status_code}: {response.text}")
    return response.json()


def find_child_folder(token: str, parent_id: str, folder_name: str):
    url = f"{GRAPH_BASE_URL}/me/mailFolders/{parent_id}/childFolders"
    while url:
        data = graph_get(token, url, params={"$top": 100})
        for folder in data.get("value", []):
            if folder.get("displayName", "").lower() == folder_name.lower():
                return folder["id"]
        url = data.get("@odata.nextLink")
    return None


def get_folder_id(token: str, folder_path: str):
    parts = [part.strip() for part in folder_path.replace("\\", "/").split("/") if part.strip()]
    if not parts:
        raise ValueError("Folder path is empty.")

    folder_id = "inbox"
    if parts[0].lower() == "inbox":
        parts = parts[1:]

    for part in parts:
        next_id = find_child_folder(token, folder_id, part)
        if not next_id:
            raise RuntimeError(f"Could not find Outlook folder: {folder_path}")
        folder_id = next_id
    return folder_id


def safe_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "_", name)
    return cleaned.strip() or "attachment"


def list_messages(token: str, folder_id: str, limit: int):
    messages = []
    url = f"{GRAPH_BASE_URL}/me/mailFolders/{folder_id}/messages"
    params = {
        "$top": min(limit, 50),
        "$orderby": "receivedDateTime desc",
        "$filter": "hasAttachments eq true",
        "$select": "id,subject,from,receivedDateTime,hasAttachments",
    }

    while url and len(messages) < limit:
        data = graph_get(token, url, params=params)
        messages.extend(data.get("value", []))
        url = data.get("@odata.nextLink")
        params = None

    return messages[:limit]


def download_cv_attachments(token: str, messages: list[dict], attachments_dir: Path):
    attachments_dir.mkdir(parents=True, exist_ok=True)
    saved = []

    for message in messages:
        attachments_url = f"{GRAPH_BASE_URL}/me/messages/{message['id']}/attachments"
        data = graph_get(token, attachments_url)
        for attachment in data.get("value", []):
            if attachment.get("@odata.type") != "#microsoft.graph.fileAttachment":
                continue

            filename = safe_filename(attachment.get("name", "attachment"))
            suffix = Path(filename).suffix.lower()
            if suffix not in {".pdf", ".docx"}:
                continue

            content = attachment.get("contentBytes")
            if not content:
                continue

            timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
            output_file = attachments_dir / f"{timestamp}_{filename}"
            output_file.write_bytes(base64.b64decode(content))
            saved.append(
                {
                    "path": output_file,
                    "subject": message.get("subject", ""),
                    "sender": message.get("from", {}).get("emailAddress", {}).get("address", ""),
                    "received_time": message.get("receivedDateTime", ""),
                }
            )

    return saved


def extract_pdf_text(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def extract_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        return extract_pdf_text(path)
    if path.suffix.lower() == ".docx":
        return extract_docx_text(path)
    return ""


def first_non_empty_line(text: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if EMAIL_PATTERN.search(line) or PHONE_PATTERN.search(line):
            continue
        if len(line.split()) <= 6:
            return line
    return ""


def extract_experience(text: str) -> str:
    matches = EXPERIENCE_PATTERN.findall(text)
    if not matches:
        return ""
    numbers = [float(match) for match in matches]
    return f"{max(numbers):g} years"


def extract_current_company(text: str) -> str:
    patterns = [
        r"(?:current\s+company|current\s+employer|present\s+company|present\s+employer)\s*[:\-]\s*(.+)",
        r"(?:working\s+with|working\s+at)\s+(.+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).splitlines()[0].strip(" .,-")
    return ""


def extract_skills(text: str) -> str:
    lower_text = text.lower()
    found = [skill for skill in KNOWN_SKILLS if skill.lower() in lower_text]
    return ", ".join(dict.fromkeys(found))


def parse_candidate_details(path: Path):
    text = extract_text(path)
    emails = sorted(set(EMAIL_PATTERN.findall(text)))
    phones = sorted(set(match.group(0).strip() for match in PHONE_PATTERN.finditer(text)))
    return {
        "name": first_non_empty_line(text),
        "email": ", ".join(emails),
        "phone": ", ".join(phones),
        "experience": extract_experience(text),
        "company": extract_current_company(text),
        "skills": extract_skills(text),
    }


def run_agent(args):
    token = get_access_token(args.client_id, args.tenant)
    folder_id = get_folder_id(token, args.folder)
    messages = list_messages(token, folder_id, args.email_limit)
    attachments = download_cv_attachments(token, messages, Path(args.attachments_dir).resolve())

    workbook = ensure_workbook(Path(args.output).resolve())
    sheet = workbook[SHEET_NAME]
    seen_files = existing_cv_files(sheet)
    processed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    added = 0

    for attachment in attachments:
        cv_path = attachment["path"]
        if str(cv_path).lower() in seen_files:
            continue

        try:
            details = parse_candidate_details(cv_path)
            sheet.append(
                [
                    details["name"],
                    details["email"],
                    details["phone"],
                    details["experience"],
                    details["company"],
                    details["skills"],
                    str(cv_path),
                    attachment["subject"],
                    attachment["sender"],
                    attachment["received_time"],
                    processed_at,
                ]
            )
            added += 1
        except Exception as exc:
            print(f"Could not process {cv_path.name}: {exc}")

    workbook.save(Path(args.output).resolve())

    if args.clean_attachments:
        shutil.rmtree(Path(args.attachments_dir).resolve(), ignore_errors=True)

    print("\nDone.")
    print(f"Emails scanned: {len(messages)}")
    print(f"CV attachments found: {len(attachments)}")
    print(f"New candidates added: {added}")
    print(f"Excel file: {Path(args.output).resolve()}")


def build_parser():
    parser = argparse.ArgumentParser(description="Read CV attachments from Outlook cloud mail using Microsoft Graph.")
    parser.add_argument("--client-id", required=True, help="Azure app registration Application client ID.")
    parser.add_argument("--tenant", default="common", help="Tenant ID, or common for work/school/personal accounts.")
    parser.add_argument("--folder", default=DEFAULT_FOLDER, help="Outlook folder name or path, for example 'CVs To Process'.")
    parser.add_argument("--output", default=DEFAULT_OUTPUT, help="Excel output file.")
    parser.add_argument("--attachments-dir", default=DEFAULT_ATTACHMENTS_DIR, help="Folder for saved CV attachments.")
    parser.add_argument("--email-limit", type=int, default=50, help="Maximum recent emails to scan from the folder.")
    parser.add_argument("--clean-attachments", action="store_true", help="Delete downloaded CV files after Excel is created.")
    return parser


if __name__ == "__main__":
    run_agent(build_parser().parse_args())
